import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import pdfplumber
import pytesseract
from PIL import Image
import docx
from transformers import pipeline
import torch

@st.cache_resource(show_spinner=True)
def load_qa_model():
    return pipeline("question-answering", model="deepset/roberta-base-squad2")

qa_pipeline = load_qa_model()

st.title("Digital Analyst App")

uploaded_file = st.file_uploader("Upload a document", type=["txt", "docx", "csv", "xlsx", "pdf", "png", "jpg"])

def process_txt(file):
    file.seek(0)
    return file.read().decode("utf-8")

def process_docx(file):
    file.seek(0)
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def process_csv(file):
    file.seek(0)
    return pd.read_csv(file)

def process_excel(file):
    file.seek(0)
    return pd.read_excel(file)

def process_pdf(file):
    file.seek(0)
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def process_image(file):
    file.seek(0)
    img = Image.open(file)
    return pytesseract.image_to_string(img)

# Additional extraction details functions:
def extract_excel_details(file):
    """
    Extract details from an Excel file using openpyxl.
    Provides sheet names and the number of rows and columns in each sheet.
    """
    file.seek(0)
    try:
        from openpyxl import load_workbook
    except ImportError:
        st.error("Missing openpyxl. Please run: pip install openpyxl")
        return None
    wb = load_workbook(filename=file, read_only=True)
    details = {}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        details[sheet] = {"rows": ws.max_row, "columns": ws.max_column}
    return details

def extract_docx_details(file):
    """
    Extract details from a Word file.
    Returns the number of paragraphs and a sample (first paragraph).
    """
    file.seek(0)
    doc = docx.Document(file)
    details = {
        "paragraph_count": len(doc.paragraphs),
        "first_paragraph": doc.paragraphs[0].text if doc.paragraphs else ""
    }
    return details

data = None
details = None
if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "txt":
        data = process_txt(uploaded_file)
    elif file_type == "docx":
        data = process_docx(uploaded_file)
        # Extract additional details for docx files
        details = extract_docx_details(uploaded_file)
    elif file_type == "csv":
        data = process_csv(uploaded_file)
    elif file_type == "xlsx":
        data = process_excel(uploaded_file)
        # Extract additional details for Excel files
        details = extract_excel_details(uploaded_file)
    elif file_type == "pdf":
        data = process_pdf(uploaded_file)
    elif file_type in ["png", "jpg"]:
        data = process_image(uploaded_file)
    
    # Visualization for DataFrame files (CSV, Excel)
    if isinstance(data, pd.DataFrame):
        st.write("## Data Preview")
        st.dataframe(data)
        
        st.write("## Data Visualization")
        chart_type = st.selectbox("Choose chart type", ["Bar Chart", "Line Chart"])
        numeric_data = data.select_dtypes(include=[np.number])
        
        if chart_type == "Bar Chart":
            st.bar_chart(numeric_data)
        elif chart_type == "Line Chart":
            st.line_chart(numeric_data)
        
        if len(numeric_data.columns) >= 2:
            if st.checkbox("Show scatter plot"):
                x_col = st.selectbox("Select X-axis", numeric_data.columns)
                y_col = st.selectbox("Select Y-axis", numeric_data.columns, index=1)
                fig, ax = plt.subplots()
                ax.scatter(numeric_data[x_col], numeric_data[y_col])
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                st.pyplot(fig)
    
    # Visualization for text-based files
    elif isinstance(data, str):
        st.write("## Extracted Text")
        st.text_area("Text Output", data, height=300)
        if st.checkbox("Visualize word frequency"):
            import collections, re
            words = re.findall(r'\w+', data.lower())
            word_counts = collections.Counter(words)
            most_common = word_counts.most_common(10)
            df_words = pd.DataFrame(most_common, columns=["Word", "Count"])
            st.bar_chart(df_words.set_index("Word"))
    
    if details:
        st.write("## File Details")
        st.json(details)



st.write("### Ask a Question About the Text")
question = st.text_input("Enter your question")

if question and data is not None:
    # Convert data to string if it's a DataFrame (i.e., CSV or Excel)
    if isinstance(data, pd.DataFrame):
        context = data.to_string()
    # Otherwise, assume it's already a string
    elif isinstance(data, str):
        context = data
    else:
        context = str(data)
    
    try:
        answer = qa_pipeline(question=question, context=context)
        st.write("### Answer:", answer["answer"])
    except Exception as e:
        st.error(f"Error during question answering: {e}")
